from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def remove_table_borders_and_shading(table):
    """Remove all colors and make table plain white with simple borders"""
    tbl = table._tbl
    for cell in tbl.iter_tcs():
        tcPr = cell.get_or_add_tcPr()
        # Remove shading (background color)
        for shd in tcPr.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd'):
            tcPr.remove(shd)

def set_table_borders(table):
    """Set simple black borders for table"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml('<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
    
    tblBorders = parse_xml(
        '<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)

def create_table(doc, headers, rows, col_widths=None):
    """Create a plain table with no colors"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for para in header_cells[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(11)
    
    # Data rows
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_data in enumerate(row_data):
            row_cells[col_idx].text = str(cell_data)
            for para in row_cells[col_idx].paragraphs:
                if col_idx == 0 or col_idx == len(row_data) - 1:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    run.font.size = Pt(10)
    
    # Set column widths if provided
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)
    
    # Remove colors and set borders
    remove_table_borders_and_shading(table)
    set_table_borders(table)
    
    return table

def main():
    doc = Document()
    
    # Set document margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    
    # ============ LIST OF FIGURES ============
    heading = doc.add_heading('List of Figures', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    figures_data = [
        ["1", "Introduction & UI Features", ""],
        ["", "1.1 Core Features Overview", "11"],
        ["", "1.2 Problem vs Solution Comparison", "12"],
        ["2", "User Interface Components", ""],
        ["", "2.1 NEXUS Match Feed UI", "13"],
        ["", "2.2 Player Profile & Portfolio", "14"],
        ["", "2.3 Discover Gamers Page", "15"],
        ["", "2.4 User Profile & Gaming Profiles", "16"],
        ["", "2.5 Custom Portfolio & Interests", "17"],
        ["", "2.6 Add Game Profile Form", "18"],
        ["3", "Infrastructure Architecture (Section 2.3)", ""],
        ["", "3.1 (9A) Client & CDN Layer (Frontend)", "19"],
        ["", "3.2 (9B) Application Layer (Backend)", "19"],
        ["", "3.3 (9C) Data & External Services Layer", "20"],
        ["4", "System Workflows", ""],
        ["", "4.1 (10A) User Journey Part 1 (Steps 1-3)", "21"],
        ["", "4.2 (10B) User Journey Part 2 (Steps 4-5)", "21"],
        ["5", "Technical Stack", ""],
        ["", "5.1 (11) Technology Stack Overview", "22"],
        ["6", "Software Architecture (Section 3.2)", ""],
        ["", "6.1 (12A) Tier 1 - Presentation Layer (Frontend)", "23"],
        ["", "6.2 (12B) Tier 2 - Application Layer (Backend)", "23"],
        ["", "6.3 (12C) Tier 3 - Data Layer (Database)", "24"],
        ["7", "Database & Real-Time Features", ""],
        ["", "7.1 (13) Database Schema (ER Diagram)", "25"],
        ["", "7.2 (14) Match Applications UI", "26"],
        ["", "7.3 (15) Voice Channels Interface", "27"],
        ["", "7.4 (16) WebSocket Communication Flow", "28"],
        ["8", "Deployment Architecture (Section 4)", ""],
        ["", "8.1 (17A) Global CDN Layer (Vercel)", "29"],
        ["", "8.2 (17B) Application Layer (Railway)", "29"],
        ["", "8.3 (17C) Data Layer (Neon PostgreSQL)", "30"],
    ]
    
    create_table(doc, ["Figure No.", "Title", "Page No."], figures_data, [3, 10, 2.5])
    
    doc.add_page_break()
    
    # ============ LIST OF TABLES ============
    heading = doc.add_heading('List of Tables', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    tables_data = [
        ["1", "Cost Analysis (MVP Phase)", "27"],
        ["2", "API Endpoints Overview", "20"],
        ["3", "Firebase SMS Pricing by Region", "32"],
        ["4", "System Performance Metrics", "25"],
        ["5", "Database Tables and Schema", "19"],
    ]
    
    create_table(doc, ["Table No.", "Title", "Page No."], tables_data, [3, 10, 2.5])
    
    doc.add_page_break()
    
    # ============ TABLE OF CONTENTS ============
    heading = doc.add_heading('Table of Contents', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    toc_data = [
        ["", "Acknowledgement", "3"],
        ["", "Abstract", "4"],
        ["", "List of Figures and Tables", "6"],
        ["1", "Introduction", "8"],
        ["", "1.1 Objectives", "9"],
        ["", "1.2 Background and Literature Survey", "10"],
        ["2", "Proposed System & Methodology", "19"],
        ["", "2.1 Problem Analysis", "19"],
        ["", "2.2 System Requirements", "19"],
        ["", "2.3 Proposed Solution Architecture", "20"],
        ["3", "System Implementation & Technical Details", "21"],
        ["", "3.1 Technical Stack", "21"],
        ["", "3.2 System Architecture", "22"],
        ["", "3.3 Database Schema", "23"],
        ["", "3.4 Key Components & Features", "24"],
        ["", "3.5 API Architecture", "25"],
        ["", "3.6 Real-Time Communication", "26"],
        ["4", "Deployment and Infrastructure", "27"],
        ["5", "Results & Discussion", "28"],
        ["6", "Conclusion & Future Works", "30"],
        ["7", "References", "32"],
        ["8", "Appendix", "33"],
    ]
    
    create_table(doc, ["S.No.", "Chapter Title", "Page No."], toc_data, [2, 11, 2.5])
    
    doc.add_page_break()
    
    # ============ APPENDIX TABLES ============
    heading = doc.add_heading('Appendix', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    # Table A1: API Endpoints
    doc.add_heading('Table A1: API Endpoints Overview', level=2)
    doc.add_paragraph()
    
    api_data = [
        ["POST", "/api/auth/register", "User registration with validation"],
        ["POST", "/api/auth/login", "Session-based authentication"],
        ["GET", "/api/auth/session", "Get current user session"],
        ["POST", "/api/auth/logout", "End user session"],
        ["GET", "/api/users", "List users with filters"],
        ["GET", "/api/users/:id", "Get user profile details"],
        ["PUT", "/api/users/:id", "Update user profile"],
        ["GET", "/api/matches", "List available matches"],
        ["POST", "/api/matches", "Create new match request"],
        ["GET", "/api/matches/:id", "Get match details"],
        ["DELETE", "/api/matches/:id", "Delete match request"],
        ["POST", "/api/matches/:id/apply", "Apply to match"],
        ["GET", "/api/connections", "List user connections"],
        ["POST", "/api/connections/request", "Send connection request"],
        ["PUT", "/api/connections/:id/accept", "Accept connection"],
        ["DELETE", "/api/connections/:id", "Remove connection"],
        ["GET", "/api/notifications", "Get user notifications"],
        ["PUT", "/api/notifications/:id/read", "Mark notification as read"],
        ["GET", "/api/games", "List supported games"],
        ["GET", "/api/game-profiles", "Get user game profiles"],
        ["POST", "/api/game-profiles", "Add game profile"],
    ]
    
    create_table(doc, ["Method", "Endpoint", "Description"], api_data, [2.5, 5, 8])
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Table A2: Database Schema
    doc.add_heading('Table A2: Database Tables and Schema', level=2)
    doc.add_paragraph()
    
    db_data = [
        ["users", "id, email, username, password_hash, avatar_url, bio, location, phone_verified, created_at"],
        ["match_requests", "id, creator_id, game_id, title, description, match_type, region, status, created_at"],
        ["user_connections", "id, requester_id, receiver_id, status, created_at, accepted_at"],
        ["notifications", "id, user_id, type, title, message, read, created_at"],
        ["games", "id, name, icon_url, category, supported_ranks"],
        ["user_game_profiles", "id, user_id, game_id, current_rank, peak_rank, hours_played, screenshot_url"],
        ["match_applications", "id, match_id, applicant_id, message, status, created_at"],
        ["voice_channels", "id, name, match_id, created_by, hms_room_id, created_at"],
        ["push_subscriptions", "id, user_id, endpoint, p256dh, auth, created_at"],
    ]
    
    create_table(doc, ["Table Name", "Columns"], db_data, [4, 12])
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Table A3: Cost Analysis
    doc.add_heading('Table A3: Cost Analysis (MVP Phase)', level=2)
    doc.add_paragraph()
    
    cost_data = [
        ["Vercel", "Frontend Hosting", "Free tier", "$0"],
        ["Railway", "Backend + Database", "Free credits ($5/mo)", "$0-5"],
        ["Neon", "PostgreSQL Database", "Free tier (3GB)", "$0"],
        ["Firebase", "Phone Auth (SMS)", "Free tier (50K/mo)", "$0"],
        ["100ms", "Voice Communication", "Free tier (10K min/mo)", "$0"],
        ["Cloudflare R2", "Object Storage", "Free tier (10GB)", "$0"],
        ["Domain", "Custom Domain", "Optional", "$0-12/yr"],
        ["", "Total Monthly Cost", "", "$0-5"],
    ]
    
    create_table(doc, ["Service", "Purpose", "Plan", "Cost"], cost_data, [3.5, 4, 4, 2.5])
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Table A4: Performance Metrics
    doc.add_heading('Table A4: System Performance Metrics', level=2)
    doc.add_paragraph()
    
    perf_data = [
        ["WebSocket Latency", "<100ms", "45ms average"],
        ["API Response Time", "<500ms", "120ms average"],
        ["Lighthouse Score", ">90", "98/100"],
        ["First Contentful Paint", "<2s", "0.8s"],
        ["Time to Interactive", "<3s", "1.2s"],
        ["Uptime", ">99%", "99.9%"],
        ["Database Query Time", "<50ms", "12ms average"],
        ["Voice Latency (100ms)", "<100ms", "45ms"],
    ]
    
    create_table(doc, ["Metric", "Target", "Achieved"], perf_data, [5, 4, 5])
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Table A5: Firebase SMS Pricing
    doc.add_heading('Table A5: Firebase SMS Pricing by Region', level=2)
    doc.add_paragraph()
    
    sms_data = [
        ["United States", "$0.01", "50,000 free/month"],
        ["India", "$0.01", "50,000 free/month"],
        ["United Kingdom", "$0.04", "50,000 free/month"],
        ["Germany", "$0.07", "50,000 free/month"],
        ["Brazil", "$0.03", "50,000 free/month"],
        ["Japan", "$0.06", "50,000 free/month"],
        ["Australia", "$0.05", "50,000 free/month"],
    ]
    
    create_table(doc, ["Region", "Cost per SMS", "Free Tier"], sms_data, [5, 4, 5])
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Table A6: Team Contributions
    doc.add_heading('Table A6: Team Member Contributions', level=2)
    doc.add_paragraph()
    
    team_data = [
        ["22BCE9357", "Adnan Hasshad Md", "Project Manager & Technical Lead", "Architecture design, Backend development, Database schema, Deployment"],
        ["22BCE9911", "Mayakuntla Lokesh", "Back-End Developer", "API development, WebSocket implementation, Authentication, Database queries"],
        ["22BCE9745", "Thokala Sravan", "Front-End Developer", "React UI development, Component design, PWA implementation, Voice integration"],
        ["22BCE20420", "Tatikonda Srilekha", "QA & Support Developer", "Testing, Documentation, Bug fixes, User support"],
    ]
    
    create_table(doc, ["Reg. No", "Name", "Role", "Contributions"], team_data, [2.5, 4, 4, 5.5])
    
    # Save document
    output_path = "attached_assets/NEXUS_TABLES_FORMATTED.docx"
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
    return output_path

if __name__ == "__main__":
    main()

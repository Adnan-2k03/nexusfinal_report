from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

def remove_empty_pages(input_path, output_path):
    doc = Document(input_path)
    
    paragraphs_to_remove = []
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        has_page_break = False
        for run in para.runs:
            if run._element.xml.find('w:br') != -1 and 'w:type="page"' in run._element.xml:
                has_page_break = True
        
        if not text and not para.runs:
            paragraphs_to_remove.append(para)
        elif not text and has_page_break:
            for run in para.runs:
                for br in run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br'):
                    br_type = br.get(qn('w:type'))
                    if br_type == 'page':
                        run._element.remove(br)
    
    for para in paragraphs_to_remove:
        p = para._element
        p.getparent().remove(p)
    
    cleaned_paragraphs = []
    prev_empty = False
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            if not prev_empty:
                prev_empty = True
            else:
                p = para._element
                p.getparent().remove(p)
        else:
            prev_empty = False
    
    doc.save(output_path)
    print(f"Cleaned document saved to: {output_path}")

input_file = "attached_assets/NEXUS_CAPSTONE_REPORT_FINAL_(1)_(1)_1764897460962.docx"
output_file = "attached_assets/NEXUS_CAPSTONE_REPORT_CLEANED.docx"

remove_empty_pages(input_file, output_file)
